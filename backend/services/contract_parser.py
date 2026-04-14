from __future__ import annotations
import re
from pathlib import Path
from typing import List
from docx import Document
from .schemas import ClauseSnippet


KEYWORDS = {
    "payment terms": [r"payment.{0,40}days", r"undisputed amounts", r"invoice"],
    "withholding and set-off": [r"withhold", r"set off|set-off", r"deduct"],
    "pod requirements": [r"POD", r"proof of delivery", r"ePOD|e-pod"],
    "loss or damage": [r"loss", r"damage", r"shortage", r"theft", r"misdelivery"],
    "audit rights": [r"audit", r"inspection", r"records"],
    "discounts / rebates": [r"discount", r"rebate", r"service credit"],
    "penalties": [r"penalt", r"sla", r"service levels?"],
}


def _read_docx_text(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def extract_contract_text(path: str) -> str:
    suffix = Path(path).suffix.lower()
    if suffix == ".docx":
        return _read_docx_text(path)
    raise ValueError(f"Unsupported contract format: {suffix}")


def extract_key_clauses(contract_text: str, limit: int = 10) -> List[ClauseSnippet]:
    snippets: List[ClauseSnippet] = []
    blocks = [b.strip() for b in re.split(r"\n+", contract_text) if b.strip()]
    for title, patterns in KEYWORDS.items():
        for block in blocks:
            for pattern in patterns:
                if re.search(pattern, block, flags=re.IGNORECASE):
                    snippets.append(ClauseSnippet(title=title.title(), snippet=block[:450]))
                    break
            if any(s.title == title.title() for s in snippets):
                break
    return snippets[:limit]


def extract_contract_metadata(contract_text: str) -> dict:
    payment_days = None
    match = re.search(r"within\s*\[?(\d{1,3})\/?(\d{0,3})?\]?\s*days", contract_text, flags=re.IGNORECASE)
    if match:
        payment_days = int(match.group(2) or match.group(1))
    return {
        "payment_days": payment_days,
        "contains_pod_requirement": bool(re.search(r"proof of delivery|\bPOD\b|ePOD|e-pod", contract_text, re.I)),
        "contains_setoff_right": bool(re.search(r"set off|set-off|withhold|deduct", contract_text, re.I)),
        "contains_damage_liability": bool(re.search(r"damage|loss|shortage|theft", contract_text, re.I)),
    }
