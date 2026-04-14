from __future__ import annotations
from pathlib import Path
from typing import Dict, Tuple
import pandas as pd
from docx import Document


AMOUNT_CANDIDATES = [
    "Net Bill Amount", "Net Amount", "Invoice Amount", "Line Amount", "Amount", "Gross Amount"
]
SHIPMENT_ID_CANDIDATES = ["Shipment ID", "shipment_id", "Shipment", "Consignment No", "AWB No"]
STATUS_CANDIDATES = ["Final Status", "Status", "Current Status"]


def _find_first_present(df: pd.DataFrame, options: list[str]) -> str | None:
    for option in options:
        if option in df.columns:
            return option
    return None


def _parse_excel_invoice(path: str) -> Tuple[dict, pd.DataFrame]:
    xls = pd.ExcelFile(path)
    detail_sheet = None
    for s in xls.sheet_names:
        if "billing" in s.lower() or "detail" in s.lower() or "annex" in s.lower():
            detail_sheet = s
            break
    if detail_sheet is None:
        detail_sheet = xls.sheet_names[0]
    df = pd.read_excel(path, sheet_name=detail_sheet)
    shipment_col = _find_first_present(df, SHIPMENT_ID_CANDIDATES)
    amount_col = _find_first_present(df, AMOUNT_CANDIDATES)
    status_col = _find_first_present(df, STATUS_CANDIDATES)
    if shipment_col is None or amount_col is None:
        raise ValueError("Could not identify shipment id and amount columns in invoice.")
    normalized = pd.DataFrame({
        "shipment_id": df[shipment_col].astype(str).str.strip(),
        "invoice_amount": pd.to_numeric(df[amount_col], errors="coerce").fillna(0.0),
        "invoice_status": df[status_col].astype(str).str.strip() if status_col else None,
    })
    normalized = normalized[normalized["shipment_id"].ne("nan") & normalized["shipment_id"].ne("")].copy()

    meta = {"invoice_number": None, "carrier_name": None, "period": None}
    if "Invoice" in xls.sheet_names:
        face = pd.read_excel(path, sheet_name="Invoice", header=None)
        flat = face.fillna("").astype(str).values.tolist()
        for row in flat:
            clean = [str(x).strip() for x in row]
            for idx, cell in enumerate(clean):
                lower = cell.lower()
                if lower == "carrier / vendor name" and idx + 1 < len(clean) and clean[idx + 1]:
                    meta["carrier_name"] = clean[idx + 1]
                elif lower in {"invoice no", "invoice number"} and idx + 1 < len(clean) and clean[idx + 1]:
                    meta["invoice_number"] = clean[idx + 1]
                elif lower == "invoice period" and idx + 1 < len(clean) and clean[idx + 1]:
                    meta["period"] = clean[idx + 1]
    return meta, normalized


def _parse_docx_invoice(path: str) -> Tuple[dict, pd.DataFrame]:
    doc = Document(path)
    meta = {"invoice_number": None, "carrier_name": None, "period": None}
    rows = []
    for p in doc.paragraphs:
        text = p.text.strip()
        lower = text.lower()
        if "invoice no" in lower:
            meta["invoice_number"] = text.split(":")[-1].strip()
        if "carrier" in lower and meta["carrier_name"] is None:
            meta["carrier_name"] = text.split(":")[-1].strip()
        if "period" in lower and meta["period"] is None:
            meta["period"] = text.split(":")[-1].strip()
    for table in doc.tables:
        header = [c.text.strip() for c in table.rows[0].cells]
        shipment_idx = next((i for i,h in enumerate(header) if h in SHIPMENT_ID_CANDIDATES), None)
        amount_idx = next((i for i,h in enumerate(header) if h in AMOUNT_CANDIDATES), None)
        if shipment_idx is None or amount_idx is None:
            continue
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            rows.append({
                "shipment_id": cells[shipment_idx],
                "invoice_amount": pd.to_numeric(cells[amount_idx], errors="coerce"),
                "invoice_status": None,
            })
    if not rows:
        raise ValueError("No bill detail table found in DOCX invoice.")
    return meta, pd.DataFrame(rows)


def load_invoice(path: str) -> Dict[str, object]:
    suffix = Path(path).suffix.lower()
    if suffix == ".xlsx":
        meta, lines = _parse_excel_invoice(path)
    elif suffix == ".docx":
        meta, lines = _parse_docx_invoice(path)
    else:
        raise ValueError("Invoice must be .xlsx or .docx for this MVP.")
    lines["shipment_id"] = lines["shipment_id"].astype(str).str.strip()
    lines["invoice_amount"] = pd.to_numeric(lines["invoice_amount"], errors="coerce").fillna(0.0)
    return {"meta": meta, "lines": lines}
